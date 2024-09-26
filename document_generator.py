import os
import sys
from docx import Document
from docx.shared import Inches
from PIL import Image
import logging
import tkinter as tk
from tkinter import messagebox

def get_safe_max_image_width(max_image_width, is_two_column):
    page_width = 8.5  # inches
    if is_two_column:
        max_allowed_width = (page_width - 2) / 2  # Account for left and right margins
        return min(max_image_width, max_allowed_width) - 0.5  # Slightly reduce to prevent cropping
    return min(max_image_width, page_width - 2)  # For one column, subtract total margins

def create_document(output_file, images_dict, max_image_width, max_image_height=4):
    doc = Document()
    folder_name = os.path.basename(output_file).replace('.docx', '')
    doc.add_heading(folder_name, level=1)

    for title, content in images_dict.items():
        doc.add_heading(title, level=2)
        max_image_width_for_layout = get_safe_max_image_width(max_image_width, content['layout'] == "Two Columns")

        if content['layout'] == "Two Columns":
            add_images_to_doc_two_columns(doc, content, max_image_width_for_layout, max_image_height)
        else:
            add_images_to_doc_one_column(doc, content, max_image_width_for_layout, max_image_height)

        if isinstance(content['note'], str) and content['note']:
            doc.add_paragraph(content['note'])
        doc.add_paragraph()

    if not save_document(doc, output_file):
        # If the user canceled the save operation, return without proceeding further
        return

def save_document(doc, output_file):
    """Attempts to save the document, prompting the user to retry or cancel if the file is open."""
    saved = False
    while not saved:
        try:
            doc.save(output_file)
            saved = True  # If the save is successful, break the loop
        except PermissionError:
            # Show a warning dialog if the file is already open
            root = tk.Tk()
            root.withdraw()  # Hide the main tkinter window
            response = messagebox.askretrycancel("Permission Denied", 
                                                 f"The file '{output_file}' is already open. Please close it and try again.")
            root.destroy()
            
            if not response:
                print("User canceled the save operation.")
                return False  # Return False to indicate that the user canceled the operation
    return True  # Return True to indicate that the save was successful

def add_images_to_doc_two_columns(doc, content, max_image_width, max_image_height):
    table = doc.add_table(rows=0, cols=2)
    for i in range(0, len(content['image_paths']), 2):
        row_cells = table.add_row().cells
        for j in range(2):
            if i + j < len(content['image_paths']):
                add_image_to_cell(row_cells[j], content['image_paths'][i + j], max_image_width, max_image_height)

def add_images_to_doc_one_column(doc, content, max_image_width, max_image_height):
    for image_path in content['image_paths']:
        add_image_to_doc(doc, image_path, max_image_width, max_image_height)

def add_image_to_cell(cell, image_path, max_image_width, max_image_height):
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            ratio = min(max_image_width / width, max_image_height / height)

            new_width = width * ratio
            new_height = height * ratio

            # Ensure the new dimensions are positive
            if new_width > 0 and new_height > 0:
                # Use the new width and height calculated from the aspect ratio
                cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(new_width), height=Inches(new_height))
    except Exception as e:
        logging.error(f"Error adding image to cell: {image_path} - {e}")

def add_image_to_doc(doc, image_path, max_image_width, max_image_height):
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            ratio = min(max_image_width / width, max_image_height / height)

            new_width = width * ratio
            new_height = height * ratio

            # Ensure the new dimensions are positive
            if new_width > 0 and new_height > 0:
                # Use the new width and height calculated from the aspect ratio
                doc.add_picture(image_path, width=Inches(new_width), height=Inches(new_height))
    except Exception as e:
        logging.error(f"Error adding image to document: {image_path} - {e}")