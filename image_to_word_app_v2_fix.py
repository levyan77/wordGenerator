import os
import json
import logging
from tkinter import Tk, Label, Button, filedialog, messagebox, Toplevel, Frame, Canvas, Scrollbar, Text, END, StringVar, OptionMenu, Entry
from docx import Document
from docx.shared import Inches
from collections import defaultdict
from PIL import Image, ImageTk
import webbrowser

# Set up logging
logging.basicConfig(filename='app.log', level=logging.ERROR)

class ImageToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Image to Word Document Generator")
        self.root.geometry("600x400")
        self.image_folder = ""
        self.notes = {}
        self.layout_choices = {}
        self.image_types = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
        self.max_image_width = 5  # Default maximum image width

        self.create_widgets()
        self.load_preferences()  # Load user preferences on start

    def create_widgets(self):
        Label(self.root, text="Select a folder containing images:", font=("Arial", 12)).pack(pady=10)
        Button(self.root, text="Select Folder", command=self.select_folder, font=("Arial", 12)).pack(pady=5)
        self.button_generate = Button(self.root, text="Generate Document", command=self.generate_document, state="disabled", font=("Arial", 12))
        self.button_generate.pack(pady=10)
        Button(self.root, text="Help", command=self.open_help, font=("Arial", 12)).pack(pady=5)
        Button(self.root, text="Feedback", command=self.open_feedback, font=("Arial", 12)).pack(pady=5)
        self.status_label = Label(self.root, text="", font=("Arial", 10))
        self.status_label.pack(pady=5)

    def update_status(self, message):
        self.status_label.config(text=message)

    def select_folder(self):
        self.image_folder = filedialog.askdirectory()
        if self.image_folder:
            self.button_generate.config(state="normal")  # Enable document generation button
            self.update_status(f"Selected folder: {os.path.basename(self.image_folder)}")
            self.show_preview()

    def show_preview(self):
        # Create a new window to preview the selected images
        self.preview_window = Toplevel(self.root)
        self.preview_window.title("Image Preview")
        self.preview_window.geometry("600x400")
        self.preview_window.resizable(True, True)

        # Set up a canvas and scrollbar for scrolling through images
        canvas = Canvas(self.preview_window)
        scrollbar = Scrollbar(self.preview_window, command=canvas.yview)
        scrollable_frame = Frame(canvas)

        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        self.load_images(scrollable_frame)  # Load and display images in the preview window

        # Entry for maximum image width and a close button
        self.max_width_entry = Entry(scrollable_frame, width=5)
        self.max_width_entry.insert(0, str(self.max_image_width))
        self.max_width_entry.pack(pady=5)
        Label(scrollable_frame, text="Max Image Width (inches)").pack(pady=5)
        Button(self.preview_window, text="Close", command=self.on_close_preview).pack(pady=10)

    def load_images(self, frame):
        images_dict = defaultdict(list)
        for filename in os.listdir(self.image_folder):
            if filename.lower().endswith(self.image_types):
                image_path = os.path.join(self.image_folder, filename)
                img = Image.open(image_path)
                img.thumbnail((100, 100))
                img_tk = ImageTk.PhotoImage(img)
                title = filename.split('_')[0]
                images_dict[title].append(image_path)

                label = Label(frame, image=img_tk)
                label.image = img_tk
                label.bind("<Button-1>", lambda e, img_path=image_path: self.show_full_image(img_path))
                label.pack(padx=5, pady=5)

        for title in images_dict.keys():
            self.add_image_controls(frame, title)

    def add_image_controls(self, frame, title):
        # Add layout options and note entry for each title
        Label(frame, text=title, font=("Arial", 14, "bold")).pack(pady=5)
        layout_var = StringVar(value="Single Column")
        OptionMenu(frame, layout_var, "Single Column", "Two Columns").pack(pady=5)
        self.layout_choices[title] = layout_var  # Store layout choice for this title

        # Create the note entry and store its reference
        note_entry = Text(frame, height=3, width=40, wrap='word')
        note_entry.pack(pady=5)
        self.notes[title] = note_entry  # Store the note entry for this title

    def show_full_image(self, image_path):
        full_image_window = Toplevel(self.preview_window)
        full_image_window.title("Full Size Image")

        self.full_image = Image.open(image_path)
        img_tk = ImageTk.PhotoImage(self.full_image)

        label = Label(full_image_window, image=img_tk)
        label.pack()

        scrollbar_x = Scrollbar(full_image_window, orient='horizontal')
        scrollbar_y = Scrollbar(full_image_window, orient='vertical')

        scrollbar_x.pack(side='bottom', fill='x')
        scrollbar_y.pack(side='right', fill='y')

        full_image_window.img = img_tk
        full_image_window.geometry(f"{self.full_image.width}x{self.full_image.height}")

        label.bind("<MouseWheel>", lambda event: self.zoom_image(event, label))

    def zoom_image(self, event, label):
        scale = 1.1 if event.delta > 0 else 0.9
        new_width = int(label.winfo_width() * scale)
        new_height = int(label.winfo_height() * scale)
        label.config(width=new_width, height=new_height)

    def on_close_preview(self):
        # Update maximum image width from user input
        try:
            self.max_image_width = float(self.max_width_entry.get())
        except ValueError:
            messagebox.showwarning("Invalid Input", "Please enter a valid number for max image width.")
            return

        # Store notes before closing the preview window
        for title in self.layout_choices.keys():
            note_widget = self.notes[title]
            if isinstance(note_widget, Text):
                self.notes[title] = note_widget.get("1.0", END).strip()  # Store the note text

        self.save_preferences()  # Save preferences
        self.preview_window.destroy()  # Close the preview window

    def generate_document(self):
        # Create the Word document from the selected images and notes
        if not self.image_folder:
            messagebox.showwarning("No Folder Selected", "Please select a folder first.")
            return

        folder_name = os.path.basename(self.image_folder)
        output_file = os.path.join(self.image_folder, f"{folder_name}.docx")

        # Check if the output file already exists
        if os.path.exists(output_file) and not messagebox.askyesno("File Exists", f"The file '{output_file}' already exists. Overwrite?"):
            return

        doc = Document()
        doc.add_heading(folder_name, level=1)  # Add main heading

        images_dict = self.compile_images()  # Compile images and associated notes

        # Define margins for the document
        left_margin = 1  # in inches
        right_margin = 1  # in inches
        total_margin = (left_margin + right_margin) * 72  # Convert margins to points

        for title, content in images_dict.items():
            doc.add_heading(title, level=2)  # Add section heading for each title
            self.add_images_to_doc(doc, content, total_margin)  # Add images to the document

            if isinstance(content['note'], str) and content['note']:  # Ensure it's a string
                doc.add_paragraph(content['note'])  # Add note below images
            doc.add_paragraph()  # Add space after each section

        doc.save(output_file)  # Save the document
        messagebox.showinfo("Success", f"Document '{output_file}' created successfully.")
        self.update_status(f"Document '{output_file}' created successfully.")

    def compile_images(self):
        images_dict = defaultdict(lambda: {'image_paths': [], 'note': '', 'layout': 'Single Column'})

        for title in self.layout_choices.keys():
            images_dict[title]['layout'] = self.layout_choices[title].get()  # Get the selected layout
            images_dict[title]['note'] = self.notes[title]  # Retrieve the stored note text

            for filename in os.listdir(self.image_folder):
                if filename.lower().endswith(self.image_types):
                    image_path = os.path.join(self.image_folder, filename)
                    if filename.split('_')[0] == title:
                        images_dict[title]['image_paths'].append(image_path)  # Collect image paths
        return images_dict

    def add_images_to_doc(self, doc, content, total_margin):
        if content['layout'] == "Two Columns":
            table = doc.add_table(rows=0, cols=2)
            for i in range(0, len(content['image_paths']), 2):
                row_cells = table.add_row().cells
                for j in range(2):
                    if i + j < len(content['image_paths']):
                        self.add_image_to_cell(row_cells[j], content['image_paths'][i + j], total_margin)
        else:
            for image_path in content['image_paths']:
                self.add_image_to_doc(doc, image_path, total_margin)

    def add_image_to_cell(self, cell, image_path, total_margin):
        with Image.open(image_path) as img:
            width, height = img.size
            max_width = (self.max_image_width * 72 / 2) - (total_margin / 2)
            ratio = min(max_width / width, 4 * 72 / height)
            new_width = width * ratio
            new_height = height * ratio
            if new_width > 0 and new_height > 0:
                cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(new_width / 72))

    def add_image_to_doc(self, doc, image_path, total_margin):
        with Image.open(image_path) as img:
            width, height = img.size
            max_width = (self.max_image_width * 72) - total_margin
            ratio = min(max_width / width, 4 * 72 / height)
            new_width = width * ratio
            new_height = height * ratio
            if new_width > 0 and new_height > 0:
                doc.add_picture(image_path, width=Inches(new_width / 72))

    def load_preferences(self):
        if os.path.exists("preferences.json"):
            with open("preferences.json", "r") as f:
                preferences = json.load(f)
                self.max_image_width = preferences.get("max_image_width", 5)

    def save_preferences(self):
        preferences = {"max_image_width": self.max_image_width}
        with open("preferences.json", "w") as f:
            json.dump(preferences, f)

    def open_help(self):
        help_window = Toplevel(self.root)
        help_window.title("Help")
        help_text = """Instructions on how to use the application:
1. Select a folder containing images.
2. Set the maximum image width in inches.
3. Click on an image to view it in full size.
4. Add notes and select layouts for each image.
5. Click 'Generate Document' to create the Word file.
"""
        Label(help_window, text=help_text, justify='left').pack(padx=10, pady=10)

    def open_feedback(self):
        webbrowser.open('https://your-feedback-url.com')  # Replace with your feedback URL

if __name__ == "__main__":
    root = Tk()
    app = ImageToWordApp(root)
    root.mainloop()