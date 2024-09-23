import os
from docx import Document
from docx.shared import Inches

def create_word_document(image_folder, output_file):
    # Create a new Document
    doc = Document()

    # Loop through the image files in the specified folder
    for filename in os.listdir(image_folder):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            # Construct full file path
            image_path = os.path.join(image_folder, filename)

            # Add title (filename without extension) to the document
            title = os.path.splitext(filename)[0]
            doc.add_heading(title, level=1)

            # Add the image to the document
            doc.add_picture(image_path, width=Inches(4.0))  # Adjust width as needed
            doc.add_paragraph()  # Add a blank line for spacing

    # Save the document
    doc.save(output_file)

if __name__ == "__main__":
    image_folder = input("Enter the path to the image folder: ")
    output_file = "output.docx"  # You can customize the output file name here

    if not os.path.exists(image_folder):
        print(f"The folder '{image_folder}' does not exist.")
    else:
        create_word_document(image_folder, output_file)
        print(f"Document '{output_file}' created successfully.")
