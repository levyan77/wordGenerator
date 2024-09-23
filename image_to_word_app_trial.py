import os
import json
import logging
from tkinter import Tk, Label, Button, filedialog, messagebox, Toplevel, Frame, Canvas, Scrollbar, Text, END, StringVar, OptionMenu, Entry
from docx import Document
from docx.shared import Inches
from collections import defaultdict
from PIL import Image, ImageTk
import webbrowser
import threading

# Set up logging
logging.basicConfig(filename='app.log', level=logging.ERROR)

class ImageToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Image to Word Document Generator")
        self.root.geometry("600x400")
        self.image_folder = ""
        self.notes = {}
        self.layout_choices = {}
        self.image_types = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
        self.max_image_width = 5  # Default maximum image width

        self.create_widgets()
        self.load_preferences()  # Load user preferences on start

    def create_widgets(self):
        Label(self.root, text="Select a folder containing images:", font=("Arial", 12)).pack(pady=10)
        Button(self.root, text="Select Folder", command=self.select_folder, font=("Arial", 12)).pack(pady=5)
        self.button_generate = Button(self.root, text="Generate Document", command=self.generate_document, state="disabled", font=("Arial", 12))
        self.button_generate.pack(pady=10)
        Button(self.root, text="Help", command=self.open_help, font=("Arial", 12)).pack(pady=5)
        Button(self.root, text="Feedback", command=self.open_feedback, font=("Arial", 12)).pack(pady=5)
        self.status_label = Label(self.root, text="", font=("Arial", 10))
        self.status_label.pack(pady=5)

    def update_status(self, message):
        self.status_label.config(text=message)

    def select_folder(self):
        self.image_folder = filedialog.askdirectory()
        if self.image_folder:
            self.button_generate.config(state="normal")  # Enable document generation button
            self.update_status(f"Selected folder: {os.path.basename(self.image_folder)}")
            self.show_preview()

    def show_preview(self):
        # Create a new window to preview the selected images
        self.preview_window = Toplevel(self.root)
        self.preview_window.title("Image Preview")
        self.preview_window.geometry("700x500")
        self.preview_window.resizable(True, True)

        # Set up a canvas and scrollbar for scrolling through images
        self.canvas = Canvas(self.preview_window)
        scrollbar = Scrollbar(self.preview_window, command=self.canvas.yview)
        self.scrollable_frame = Frame(self.canvas)

        self.scrollable_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # Title for the preview
        Label(self.scrollable_frame, text="Image Preview", font=("Arial", 16, "bold")).pack(pady=10)

        self.load_images(self.scrollable_frame)  # Load and display images in the preview window

        # Entry for maximum image width
        width_frame = Frame(self.scrollable_frame)
        width_frame.pack(pady=15)

        Label(width_frame, text="Max Image Width (inches):", font=("Arial", 12)).pack(side="left")
        self.max_width_entry = Entry(width_frame, width=5, font=("Arial", 12))
        self.max_width_entry.insert(0, str(self.max_image_width))
        self.max_width_entry.pack(side="left", padx=5)

        # Save/Apply button
        Button(self.preview_window, text="Save/Apply", command=self.on_save_apply, font=("Arial", 12), bg="lightgreen").pack(pady=10)

    def load_images(self, frame):
        images_dict = defaultdict(list)
        for filename in os.listdir(self.image_folder):
            if filename.lower().endswith(self.image_types):
                image_path = os.path.join(self.image_folder, filename)
                try:
                    title = filename.split('_')[0]
                    images_dict[title].append(image_path)
                except Exception as e:
                    logging.error(f"Failed to load image {image_path}: {e}")
                    messagebox.showwarning("Load Error", f"Could not load image: {filename}. Skipping.")
        
        # Add image controls and preview
        for title, image_paths in images_dict.items():
            self.add_image_controls(frame, title, image_paths)

    def add_image_controls(self, frame, title, image_paths):
        # Add layout options and note entry for each title
        title_frame = Frame(frame)
        title_frame.pack(pady=10)

        Label(title_frame, text=title, font=("Arial", 14, "bold")).pack(pady=5)

        layout_var = StringVar(value="Single Column")
        # Bind the refresh_preview function with a lambda to pass the title
        option_menu = OptionMenu(title_frame, layout_var, "Single Column", "Two Columns", 
                                command=lambda _: self.refresh_preview(title))  
        option_menu.pack(side="left", padx=5)
        self.layout_choices[title] = layout_var  # Store layout choice for this title

        # Create the note entry and store its reference
        note_entry = Text(title_frame, height=3, width=40, wrap='word', font=("Arial", 12))
        note_entry.pack(pady=5)
        self.notes[title] = note_entry  # Store the note entry for this title

        # Initial preview of images for this title
        self.preview_images(title_frame, title, image_paths)


    def preview_images(self, frame, title, image_paths):
        """Display images in a given layout (Single or Two Columns)."""
        layout = self.layout_choices[title].get()
        columns = 2 if layout == "Two Columns" else 1

        preview_frame = Frame(frame)
        preview_frame.pack(pady=5)

        for i, image_path in enumerate(image_paths):
            img = Image.open(image_path)
            img.thumbnail((100, 100))
            img_tk = ImageTk.PhotoImage(img)

            label = Label(preview_frame, image=img_tk)
            label.image = img_tk
            label.grid(row=i // columns, column=i % columns, padx=5, pady=5)

            label.bind("<Button-1>", lambda e, img_path=image_path: self.show_full_image(img_path))

        # Ensure that the preview frame is placed in the correct location
        frame.pack(pady=10)


    def refresh_preview(self, title):
        """Refresh the image previews for the specified title when layout changes."""
        # Find the corresponding frame and remove only the image preview section
        for widget in self.scrollable_frame.winfo_children():
            if isinstance(widget, Frame) and widget.winfo_children()[0].cget("text") == title:
                # Remove only the image preview (not the whole title section)
                for subwidget in widget.winfo_children():
                    if isinstance(subwidget, Frame):  # Assuming the image previews are wrapped in a Frame
                        subwidget.destroy()

                # Reload images with updated layout
                image_paths = [os.path.join(self.image_folder, filename) for filename in os.listdir(self.image_folder) if filename.startswith(title)]
                self.preview_images(widget, title, image_paths)
                break

    def show_full_image(self, image_path):
        full_image_window = Toplevel(self.preview_window)
        full_image_window.title("Full Size Image")

        self.full_image = Image.open(image_path)
        img_tk = ImageTk.PhotoImage(self.full_image)

        label = Label(full_image_window, image=img_tk)
        label.pack()

        scrollbar_x = Scrollbar(full_image_window, orient='horizontal')
        scrollbar_y = Scrollbar(full_image_window, orient='vertical')

        scrollbar_x.pack(side='bottom', fill='x')
        scrollbar_y.pack(side='right', fill='y')

        full_image_window.img = img_tk
        full_image_window.geometry(f"{self.full_image.width}x{self.full_image.height}")

        label.bind("<MouseWheel>", lambda event: self.zoom_image(event, label))

    def zoom_image(self, event, label):
        scale = 1.1 if event.delta > 0 else 0.9
        new_width = int(label.winfo_width() * scale)
        new_height = int(label.winfo_height() * scale)
        label.config(width=new_width, height=new_height)

    def on_save_apply(self):
        # Update maximum image width from user input
        try:
            self.max_image_width = float(self.max_width_entry.get())
        except ValueError:
            messagebox.showwarning("Invalid Input", "Please enter a valid number for max image width.")
            return

        # Store notes before closing the preview window
        for title in self.layout_choices.keys():
            note_widget = self.notes[title]
            if isinstance(note_widget, Text):
                self.notes[title] = note_widget.get("1.0", END).strip()  # Store the note text

        self.save_preferences()  # Save preferences
        messagebox.showinfo("Success", "Settings have been saved!")  # Notify the user
        self.preview_window.destroy()  # Close the preview window

    def generate_document(self):
        self.update_status("Generating document...")
        threading.Thread(target=self.create_document).start()  # Run in a separate thread

    def create_document(self):
        # Create the Word document from the selected images and notes
        if not self.image_folder:
            messagebox.showwarning("No Folder Selected", "Please select a folder first.")
            return

        folder_name = os.path.basename(self.image_folder)
        output_file = os.path.join(self.image_folder, f"{folder_name}.docx")

        # Check if the output file already exists
        if os.path.exists(output_file) and not messagebox.askyesno("File Exists", f"The file '{output_file}' already exists. Overwrite?"):
            return

        doc = Document()
        doc.add_heading(folder_name, level=1)  # Add main heading

        images_dict = self.compile_images()  # Compile images and associated notes

        # Define margins for the document
        left_margin = 1  # in inches
        right_margin = 1  # in inches
        total_margin = (left_margin + right_margin) * 72  # Convert margins to points

        for title, content in images_dict.items():
            doc.add_heading(title, level=2)  # Add section heading for each title
            self.add_images_to_doc(doc, content, total_margin)  # Add images to the document

            if isinstance(content['note'], str) and content['note']:  # Ensure it's a string
                doc.add_paragraph(content['note'])  # Add note below images
            doc.add_paragraph()  # Add space after each section

        doc.save(output_file)  # Save the document
        messagebox.showinfo("Success", f"Document '{output_file}' created successfully.")
        self.update_status(f"Document '{output_file}' created successfully.")

    def compile_images(self):
        images_dict = defaultdict(lambda: {'image_paths': [], 'note': '', 'layout': 'Single Column'})

        for title in self.layout_choices.keys():
            images_dict[title]['layout'] = self.layout_choices[title].get()  # Get the selected layout
            images_dict[title]['note'] = self.notes[title]  # Retrieve the stored note text

            for filename in os.listdir(self.image_folder):
                if filename.lower().endswith(self.image_types):
                    image_path = os.path.join(self.image_folder, filename)
                    if filename.split('_')[0] == title:
                        images_dict[title]['image_paths'].append(image_path)  # Collect image paths
        return images_dict

    def add_images_to_doc(self, doc, content, total_margin):
        if content['layout'] == "Two Columns":
            table = doc.add_table(rows=0, cols=2)
            for i in range(0, len(content['image_paths']), 2):
                row_cells = table.add_row().cells
                for j in range(2):
                    if i + j < len(content['image_paths']):
                        self.add_image_to_cell(row_cells[j], content['image_paths'][i + j], total_margin)
        else:
            for image_path in content['image_paths']:
                self.add_image_to_doc(doc, image_path, total_margin)

    def add_image_to_cell(self, cell, image_path, total_margin):
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                max_width = (self.max_image_width * 72 / 2) - (total_margin / 2)
                ratio = min(max_width / width, 4 * 72 / height)
                new_width = width * ratio
                new_height = height * ratio
                if new_width > 0 and new_height > 0:
                    cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(new_width / 72))
        except Exception as e:
            logging.error(f"Error adding image to document: {image_path} - {e}")

    def add_image_to_doc(self, doc, image_path, total_margin):
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                max_width = (self.max_image_width * 72) - total_margin
                ratio = min(max_width / width, 4 * 72 / height)
                new_width = width * ratio
                new_height = height * ratio
                if new_width > 0 and new_height > 0:
                    doc.add_picture(image_path, width=Inches(new_width / 72))
        except Exception as e:
            logging.error(f"Error adding image to document: {image_path} - {e}")


    def open_help(self):
        help_url = "https://www.example.com/help"
        webbrowser.open(help_url)

    def open_feedback(self):
        feedback_url = "https://www.example.com/feedback"
        webbrowser.open(feedback_url)

    def save_preferences(self):
        preferences = {"max_image_width": self.max_image_width}
        with open('preferences.json', 'w') as pref_file:
            json.dump(preferences, pref_file)

    def load_preferences(self):
        if os.path.exists('preferences.json'):
            with open('preferences.json', 'r') as pref_file:
                preferences = json.load(pref_file)
                self.max_image_width = preferences.get("max_image_width", self.max_image_width)

if __name__ == "__main__":
    root = Tk()
    app = ImageToWordApp(root)
    root.mainloop()